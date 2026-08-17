"""
DOCX Resume Parser

Extracts structured data from Microsoft Word (.docx) resume files.
Uses python-docx for parsing (free, no API costs).

Author: Free Hall <whall4.wh@gmail.com>
Organization: 7 Eagle Group
"""

from pathlib import Path
from typing import Union

try:
    from docx import Document
except ImportError:
    raise ImportError(
        "python-docx is required for DOCX parsing. " "Install with: pip install python-docx"
    )

from .base_parser import BaseResumeParser
from .schema import ResumeSchema


class DOCXResumeParser(BaseResumeParser):
    """Parser for Microsoft Word (.docx) resume files"""

    def __init__(self):
        super().__init__()
        self.supported_formats = [".docx"]

    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract all text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Raw text content
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not self.is_supported(path):
            raise ValueError(f"Unsupported file format: {path.suffix}")

        try:
            doc = Document(str(path))
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            return self.clean_text(full_text)

        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX text: {str(e)}")

    def parse(self, file_path: Union[str, Path]) -> ResumeSchema:
        """
        Parse DOCX resume into structured schema

        Args:
            file_path: Path to DOCX resume file

        Returns:
            ResumeSchema with extracted data
        """
        raw_text = self.extract_text(file_path)

        # Reuse PDF parser logic (both work on text)
        from .pdf_parser import PDFResumeParser

        pdf_parser = PDFResumeParser()

        # Use PDF parser methods for text analysis
        full_name = pdf_parser._extract_name(raw_text)
        email = self.extract_email(raw_text)
        phone = self.extract_phone(raw_text)
        location = pdf_parser._extract_location(raw_text)
        linkedin = self.extract_urls(raw_text, "linkedin.com")
        github = self.extract_urls(raw_text, "github.com")

        skills = pdf_parser._extract_skills(raw_text)
        experience = pdf_parser._extract_experience(raw_text)
        education = pdf_parser._extract_education(raw_text)
        certifications = pdf_parser._extract_certifications(raw_text)
        military_info = pdf_parser._extract_military_info(raw_text)

        return ResumeSchema(
            full_name=full_name,
            email=email,
            phone=phone,
            location=location,
            linkedin_url=linkedin,
            github_url=github,
            skills=skills,
            experience=experience,
            education=education,
            certifications=certifications,
            military_branch=military_info.get("branch"),
            military_mos=military_info.get("mos"),
            security_clearance=military_info.get("clearance"),
            years_of_service=military_info.get("years"),
            raw_text=raw_text,
            parse_timestamp=self.get_timestamp(),
        )

    def get_document_properties(self, file_path: Union[str, Path]) -> dict:
        """
        Extract document metadata/properties

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary of document properties
        """
        path = Path(file_path)
        doc = Document(str(path))

        props = doc.core_properties
        return {
            "author": props.author,
            "created": props.created,
            "modified": props.modified,
            "title": props.title,
            "subject": props.subject,
            "keywords": props.keywords,
        }
